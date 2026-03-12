#!/usr/bin/env python3
"""Convert Markdown to DOCX with rich formatting.

A self-contained Markdown-to-DOCX converter using regex-based parsing
and python-docx. Handles headings, bold/italic, links, code blocks,
blockquotes, lists, smart punctuation, and horizontal rules.

Usage:
    python md_to_docx.py input.md [output.docx]

If output is not specified, uses the input stem + .docx.
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Smart punctuation
# ---------------------------------------------------------------------------

def smart_punctuation(text: str) -> str:
    """Convert straight quotes to curly, -- to em dash, ... to ellipsis."""
    # Em dash (must come before en dash to avoid partial match)
    text = text.replace("---", "\u2014")
    text = text.replace("--", "\u2014")
    # Ellipsis
    text = text.replace("...", "\u2026")
    # Curly double quotes
    text = re.sub(r'"(\S)', '\u201c\\1', text)        # opening "
    text = re.sub(r'(\S)"', '\\1\u201d', text)        # closing "
    text = re.sub(r'"', '\u201d', text)                # fallback closing
    # Curly single quotes / apostrophes
    text = re.sub(r"'(\S)", '\u2018\\1', text)         # opening '
    text = re.sub(r"(\S)'", '\\1\u2019', text)         # closing / apostrophe
    text = re.sub(r"'", '\u2019', text)                 # fallback
    return text


# ---------------------------------------------------------------------------
# Inline formatting parser
# ---------------------------------------------------------------------------

# Pattern order matters: bold+italic first, then bold, then italic, then code,
# then links. Each token is (type, text, url_or_None).
INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*)'          # bold+italic
    r'|(\*\*(.+?)\*\*)'             # bold
    r'|(\*(.+?)\*)'                 # italic
    r'|(`([^`]+?)`)'                # inline code
    r'|(\[([^\]]+)\]\(([^)]+)\))'   # link
)


def _parse_inline(text: str):
    """Yield (style, content, url) tuples for inline formatting."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            yield ("plain", text[pos:m.start()], None)
        if m.group(2) is not None:      # bold+italic
            yield ("bold_italic", m.group(2), None)
        elif m.group(4) is not None:     # bold
            yield ("bold", m.group(4), None)
        elif m.group(6) is not None:     # italic
            yield ("italic", m.group(6), None)
        elif m.group(8) is not None:     # inline code
            yield ("code", m.group(8), None)
        elif m.group(10) is not None:    # link
            yield ("link", m.group(10), m.group(11))
        pos = m.end()
    # Trailing plain text
    if pos < len(text):
        yield ("plain", text[pos:], None)


def add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run_elem.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def write_inline(paragraph, text: str):
    """Parse inline markdown and add formatted runs to a paragraph."""
    text = smart_punctuation(text)
    for style, content, url in _parse_inline(text):
        if style == "plain":
            run = paragraph.add_run(content)
        elif style == "bold":
            run = paragraph.add_run(content)
            run.bold = True
        elif style == "italic":
            run = paragraph.add_run(content)
            run.italic = True
        elif style == "bold_italic":
            run = paragraph.add_run(content)
            run.bold = True
            run.italic = True
        elif style == "code":
            run = paragraph.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif style == "link":
            add_hyperlink(paragraph, content, url)


# ---------------------------------------------------------------------------
# Block-level parser
# ---------------------------------------------------------------------------

HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)')
HR_RE = re.compile(r'^---+\s*$')
FENCE_RE = re.compile(r'^```')
BLOCKQUOTE_RE = re.compile(r'^>\s?(.*)')
UL_RE = re.compile(r'^(\s*)[-*+]\s+(.*)')
OL_RE = re.compile(r'^(\s*)\d+\.\s+(.*)')


def convert(md_text: str, output_path: str) -> None:
    """Convert markdown text to a DOCX file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # --- Fenced code block ---
        if FENCE_RE.match(line):
            code_lines = []
            i += 1
            while i < len(lines) and not FENCE_RE.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            # Skip closing fence
            if i < len(lines):
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            continue

        # --- Horizontal rule ---
        if HR_RE.match(line):
            # Add a thin horizontal line via a bottom border on an empty paragraph
            para = doc.add_paragraph()
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # --- Heading ---
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            para = doc.add_heading(level=level)
            write_inline(para, heading_text)
            i += 1
            continue

        # --- Blockquote ---
        m = BLOCKQUOTE_RE.match(line)
        if m:
            quote_lines = [m.group(1)]
            i += 1
            while i < len(lines):
                bq = BLOCKQUOTE_RE.match(lines[i])
                if bq:
                    quote_lines.append(bq.group(1))
                    i += 1
                else:
                    break
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            # Italic style for blockquotes
            write_inline(para, " ".join(quote_lines))
            for run in para.runs:
                run.italic = True
            continue

        # --- Unordered list ---
        m = UL_RE.match(line)
        if m:
            indent_level = len(m.group(1)) // 2
            para = doc.add_paragraph(style="List Bullet")
            if indent_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * (indent_level + 1))
            write_inline(para, m.group(2))
            i += 1
            continue

        # --- Ordered list ---
        m = OL_RE.match(line)
        if m:
            indent_level = len(m.group(1)) // 2
            para = doc.add_paragraph(style="List Number")
            if indent_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * (indent_level + 1))
            write_inline(para, m.group(2))
            i += 1
            continue

        # --- Blank line (paragraph separator) ---
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph ---
        # Collect continuation lines (non-blank, non-special)
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if (not next_line.strip()
                    or HEADING_RE.match(next_line)
                    or HR_RE.match(next_line)
                    or FENCE_RE.match(next_line)
                    or BLOCKQUOTE_RE.match(next_line)
                    or UL_RE.match(next_line)
                    or OL_RE.match(next_line)):
                break
            para_lines.append(next_line)
            i += 1

        para = doc.add_paragraph()
        write_inline(para, " ".join(para_lines))

    doc.save(output_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert Markdown to DOCX with rich formatting."
    )
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument(
        "output",
        nargs="?",
        default=None,
        help="Output DOCX file (default: input stem + .docx)",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or str(input_path.with_suffix(".docx"))

    md_text = input_path.read_text(encoding="utf-8")
    convert(md_text, output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
