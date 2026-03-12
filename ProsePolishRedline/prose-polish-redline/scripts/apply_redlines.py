#!/usr/bin/env python3
"""Apply a review JSON to a .docx document, producing a redlined document.

Self-contained script — no local imports required. All redline logic,
summary generation, and JSON extraction are inlined.

Usage:
    python scripts/apply_redlines.py <input.docx> <review.json> [output_dir]

Outputs:
    {stem}_reviewed.docx    — redlined document with tracked changes + comments
    {stem}_review_summary.md — markdown summary with tier/severity counts
"""
import json
import logging
import re
import sys
from copy import deepcopy
from datetime import datetime, timezone
from itertools import count
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================
# TIER ORDER — Prose Polish review tiers
# ============================================================
TIER_ORDER = ["STRUCTURAL", "COHERENCE", "AUTHORITY", "CRAFT", "VOICE"]


# ============================================================
# JSON EXTRACTION (from review_engine.py)
# ============================================================

def extract_json(text: str) -> str:
    """Extract JSON from response that may contain markdown fences or commentary."""
    # Try 1: Strip markdown fences
    fence_match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", text, re.DOTALL)
    if fence_match:
        return fence_match.group(1).strip()

    # Try 2: Find JSON object boundaries
    first_brace = text.find("{")
    if first_brace != -1:
        depth = 0
        for i in range(first_brace, len(text)):
            if text[i] == "{":
                depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    return text[first_brace : i + 1]

    # Try 3: Return as-is
    return text


# ============================================================
# TRACKED CHANGE XML BUILDERS (from docx_redline.py)
# ============================================================

# Monotonic counter for unique w:id attributes (ECMA-376 requirement)
_revision_id_counter = count(1)

# Smart quote + dash mapping for fault-tolerant matching.
_QUOTE_MAP = str.maketrans({
    "\u2018": "'",   # LEFT SINGLE QUOTATION MARK
    "\u2019": "'",   # RIGHT SINGLE QUOTATION MARK
    "\u201c": '"',   # LEFT DOUBLE QUOTATION MARK
    "\u201d": '"',   # RIGHT DOUBLE QUOTATION MARK
    "\u2014": "-",   # EM DASH
    "\u2013": "-",   # EN DASH
    "\u2012": "-",   # FIGURE DASH
    "\u2015": "-",   # HORIZONTAL BAR
})


def _normalize(text: str) -> str:
    """Normalize text for matching: strip tabs, smart quotes, and dashes."""
    return text.replace("\t", "").translate(_QUOTE_MAP)


def _make_ins(text: str, author: str) -> OxmlElement:
    """Create a <w:ins> element (tracked insertion)."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(next(_revision_id_counter)))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _make_del(text: str, author: str) -> OxmlElement:
    """Create a <w:del> element (tracked deletion)."""
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(next(_revision_id_counter)))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    d.append(r)
    return d


def _make_text_run(text: str, template_run=None) -> OxmlElement:
    """Create a w:r element with text, optionally copying rPr from template."""
    r = OxmlElement("w:r")
    if template_run is not None:
        rpr = template_run.find(qn("w:rPr"))
        if rpr is not None:
            r.append(deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _set_run_text(run_elem, text: str) -> None:
    """Set the text of a w:r element, replacing all w:t children."""
    for t in run_elem.findall(qn("w:t")):
        run_elem.remove(t)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_elem.append(t)


def _flatten_hyperlinks(para_elem) -> None:
    """Move runs from w:hyperlink elements to be direct paragraph children.

    Citation references (e.g., [27]) live inside w:hyperlink wrappers.
    Flattening makes them visible to the run-level matching algorithm.
    """
    for child in list(para_elem):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "hyperlink":
            for sub in reversed(list(child)):
                child.addnext(sub)
            para_elem.remove(child)


# ============================================================
# CROSS-RUN TEXT MATCHING (from docx_redline.py)
# ============================================================

def _find_and_replace_across_runs(para, old: str, new: str, author: str) -> bool:
    """Find `old` text across potentially split runs and replace with tracked change."""
    para_elem = para._element
    _flatten_hyperlinks(para_elem)

    runs_info = []
    pos = 0
    for child in list(para_elem):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            run_text = ""
            for t in child.findall(qn("w:t")):
                run_text += t.text or ""
            runs_info.append((child, run_text, pos))
            pos += len(run_text)

    full_text = "".join(info[1] for info in runs_info)
    match_start = _normalize(full_text).find(old)
    if match_start == -1:
        return False
    match_end = match_start + len(old)

    new_elements = []
    for run_elem, run_text, run_start in runs_info:
        run_end = run_start + len(run_text)

        if run_end <= match_start or run_start >= match_end:
            new_elements.append(("keep", run_elem))
        elif run_start >= match_start and run_end <= match_end:
            new_elements.append(("delete", run_elem, run_text))
        elif run_start < match_start and run_end > match_end:
            offset_start = match_start - run_start
            offset_end = match_end - run_start
            before = run_text[:offset_start]
            matched = run_text[offset_start:offset_end]
            after = run_text[offset_end:]
            new_elements.append(("split_3", run_elem, before, matched, after))
        elif run_start < match_start:
            offset = match_start - run_start
            before = run_text[:offset]
            matched = run_text[offset:]
            new_elements.append(("split_start", run_elem, before, matched))
        elif run_end > match_end:
            offset = match_end - run_start
            matched = run_text[:offset]
            after = run_text[offset:]
            new_elements.append(("split_end", run_elem, matched, after))

    del_elem = _make_del(old, author)
    ins_elem = _make_ins(new, author)
    inserted_change = False

    for item in reversed(new_elements):
        action = item[0]
        run_elem = item[1]

        if action == "keep":
            continue
        elif action == "delete":
            if not inserted_change:
                run_elem.addnext(ins_elem)
                run_elem.addnext(del_elem)
                inserted_change = True
            para_elem.remove(run_elem)
        elif action == "split_3":
            _, _, before, matched, after = item
            if after:
                after_run = _make_text_run(after, run_elem)
                run_elem.addnext(after_run)
            run_elem.addnext(ins_elem)
            run_elem.addnext(del_elem)
            inserted_change = True
            if before:
                _set_run_text(run_elem, before)
            else:
                para_elem.remove(run_elem)
        elif action == "split_start":
            _, _, before, matched = item
            if not inserted_change:
                run_elem.addnext(ins_elem)
                run_elem.addnext(del_elem)
                inserted_change = True
            if before:
                _set_run_text(run_elem, before)
            else:
                para_elem.remove(run_elem)
        elif action == "split_end":
            _, _, matched, after = item
            if after:
                after_run = _make_text_run(after, run_elem)
                run_elem.addnext(after_run)
            if not inserted_change:
                run_elem.addnext(ins_elem)
                run_elem.addnext(del_elem)
                inserted_change = True
            para_elem.remove(run_elem)

    return True


# ============================================================
# PUBLIC TRACKED CHANGE FUNCTIONS (from docx_redline.py)
# ============================================================

def apply_tracked_replacement(
    doc: Document, old: str, new: str, author: str = "Reviewer"
) -> bool:
    """Replace first occurrence of `old` with `new` as a tracked change."""
    old = _normalize(old)
    for para in doc.paragraphs:
        if old not in _normalize(para.text):
            continue
        if _find_and_replace_across_runs(para, old, new, author):
            return True
    return False


def apply_tracked_deletion(
    doc: Document, text: str, author: str = "Reviewer"
) -> bool:
    """Delete first occurrence of `text` as a tracked change."""
    text = _normalize(text)
    for para in doc.paragraphs:
        if text not in _normalize(para.text):
            continue
        para_elem = para._element
        _flatten_hyperlinks(para_elem)

        runs_info = []
        pos = 0
        for child in list(para_elem):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "r":
                run_text = ""
                for t in child.findall(qn("w:t")):
                    run_text += t.text or ""
                runs_info.append((child, run_text, pos))
                pos += len(run_text)

        full_text = "".join(info[1] for info in runs_info)
        match_start = _normalize(full_text).find(text)
        if match_start == -1:
            continue
        match_end = match_start + len(text)

        del_elem = _make_del(text, author)
        inserted_del = False

        for run_elem, run_text, run_start in runs_info:
            run_end = run_start + len(run_text)

            if run_end <= match_start or run_start >= match_end:
                continue
            elif run_start >= match_start and run_end <= match_end:
                if not inserted_del:
                    run_elem.addnext(del_elem)
                    inserted_del = True
                para_elem.remove(run_elem)
            elif run_start < match_start and run_end > match_end:
                offset_start = match_start - run_start
                offset_end = match_end - run_start
                before = run_text[:offset_start]
                after = run_text[offset_end:]
                if after:
                    after_run = _make_text_run(after, run_elem)
                    run_elem.addnext(after_run)
                run_elem.addnext(del_elem)
                inserted_del = True
                if before:
                    _set_run_text(run_elem, before)
                else:
                    para_elem.remove(run_elem)
            elif run_start < match_start:
                offset = match_start - run_start
                before = run_text[:offset]
                if not inserted_del:
                    run_elem.addnext(del_elem)
                    inserted_del = True
                _set_run_text(run_elem, before)
            elif run_end > match_end:
                offset = match_end - run_start
                after = run_text[offset:]
                if not inserted_del:
                    run_elem.addnext(del_elem)
                    inserted_del = True
                _set_run_text(run_elem, after)

        if inserted_del:
            return True
    return False


def apply_tracked_insertion(
    doc: Document, after: str, text: str, author: str = "Reviewer"
) -> bool:
    """Insert `text` after first occurrence of `after` as a tracked change."""
    after = _normalize(after)
    for para in doc.paragraphs:
        if after not in _normalize(para.text):
            continue
        para_elem = para._element
        _flatten_hyperlinks(para_elem)

        runs_info = []
        pos = 0
        for child in list(para_elem):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "r":
                run_text = ""
                for t in child.findall(qn("w:t")):
                    run_text += t.text or ""
                runs_info.append((child, run_text, pos))
                pos += len(run_text)

        full_text = "".join(info[1] for info in runs_info)
        match_start = _normalize(full_text).find(after)
        if match_start == -1:
            continue
        match_end = match_start + len(after)

        for run_elem, run_text, run_start in runs_info:
            run_end = run_start + len(run_text)
            if run_end < match_end:
                continue
            if run_start >= match_end:
                continue

            offset_in_run = match_end - run_start
            if offset_in_run < len(run_text):
                before_text = run_text[:offset_in_run]
                after_text = run_text[offset_in_run:]
                after_run = _make_text_run(after_text, run_elem)
                run_elem.addnext(after_run)
                run_elem.addnext(_make_ins(text, author))
                _set_run_text(run_elem, before_text)
            else:
                run_elem.addnext(_make_ins(text, author))
            return True
    return False


def add_comment(
    doc: Document,
    paragraph_index: int,
    text: str,
    author: str = "Reviewer",
    initials: str = "R",
) -> None:
    """Add a margin comment to a paragraph."""
    para = doc.paragraphs[paragraph_index]
    runs = para.runs if para.runs else [para.add_run("")]
    doc.add_comment(
        runs=runs,
        text=text,
        author=author,
        initials=initials,
    )


# ============================================================
# RedlineDoc CLASS (from docx_redline.py)
# ============================================================

class RedlineDoc:
    """High-level API for building a redlined document with match tracking."""

    def __init__(self, doc: Document, author: str = "Prose Polish"):
        self.doc = doc
        self.author = author
        self.match_count = 0
        self.miss_count = 0
        self.comment_count = 0

    def replace(self, old: str, new: str) -> bool:
        result = apply_tracked_replacement(self.doc, old, new, self.author)
        if result:
            self.match_count += 1
        else:
            self.miss_count += 1
        return result

    def delete(self, text: str) -> bool:
        result = apply_tracked_deletion(self.doc, text, self.author)
        if result:
            self.match_count += 1
        else:
            self.miss_count += 1
        return result

    def insert(self, after: str, text: str) -> bool:
        result = apply_tracked_insertion(self.doc, after, text, self.author)
        if result:
            self.match_count += 1
        else:
            self.miss_count += 1
        return result

    def comment(
        self, paragraph_index: int, text: str, severity: str = "SUGGESTION"
    ) -> None:
        tagged = f"[{severity}] {text}" if severity else text
        add_comment(self.doc, paragraph_index, tagged, self.author)

    @property
    def match_rate(self) -> float:
        total = self.match_count + self.miss_count
        return self.match_count / total if total > 0 else 0.0

    @property
    def total_edits(self) -> int:
        return self.match_count + self.miss_count + self.comment_count

    def save(self, path: str | Path) -> None:
        self.doc.save(str(path))


# ============================================================
# SUMMARY GENERATION (from summary_writer.py)
# ============================================================

def generate_summary(
    edits: list[dict],
    summary: str,
    audience: str,
    match_rate: float | None = None,
    match_count: int = 0,
    miss_count: int = 0,
    comment_count: int = 0,
) -> str:
    """Generate a markdown review summary."""
    tier_counts: dict[str, int] = {}
    severity_counts: dict[str, int] = {}
    for edit in edits:
        tier = edit.get("tier", "UNKNOWN")
        tier_counts[tier] = tier_counts.get(tier, 0) + 1
        sev = edit.get("severity", "UNKNOWN")
        severity_counts[sev] = severity_counts.get(sev, 0) + 1

    lines = [
        "# Prose Polish Review Summary",
        "",
        f"**Audience:** {audience.upper()}",
        f"**Total Issues:** {len(edits)}",
    ]

    if match_rate is not None:
        parts = [f"{match_count} applied"]
        if comment_count > 0:
            parts.append(f"{comment_count} comment-only")
        parts.append(f"{miss_count} not found in document")
        lines.append(f"**Edit Match Rate:** {match_rate:.0%} ({', '.join(parts)})")
        if match_rate < 0.8:
            lines.append("")
            lines.append("> **Warning:** Match rate below 80%. Some edits could not be located in the document text.")

    lines.extend(["", "## Issues by Tier", ""])
    for tier, cnt in sorted(tier_counts.items()):
        lines.append(f"- **{tier}:** {cnt}")

    lines.extend(["", "## Issues by Severity", ""])
    for sev, cnt in sorted(severity_counts.items()):
        lines.append(f"- **{sev}:** {cnt}")

    lines.extend(["", "## Overall Assessment", "", summary, ""])

    must_fix = [e for e in edits if e.get("severity") == "MUST FIX"]
    if must_fix:
        lines.extend(["## Priority Issues (Must Fix)", ""])
        for i, edit in enumerate(must_fix, 1):
            lines.append(f"{i}. **[{edit.get('tier', '')}]** {edit.get('comment', '')}")
        lines.append("")

    lines.extend([
        "---",
        "",
        "*Generated by Prose Polish Redline System.*",
    ])

    return "\n".join(lines)


# ============================================================
# MAIN: APPLY REVIEW
# ============================================================

def apply_review(
    input_path: Path,
    review_json_path: Path,
    output_dir: Path,
    author: str = "Prose Polish",
) -> dict:
    """Apply review edits to a .docx document."""
    # Load and parse review JSON
    raw = review_json_path.read_text()
    clean = extract_json(raw)
    data = json.loads(clean)
    edits = data.get("edits", [])
    summary = data.get("summary", "")
    audience = data.get("audience", "general")

    print(f"Loaded {len(edits)} edits from {review_json_path.name}")

    doc = Document(str(input_path))
    rdoc = RedlineDoc(doc, author=author)

    for edit in edits:
        edit_type = edit.get("edit_type", "comment")
        original_text = edit.get("original_text", "")
        new_text = edit.get("new_text", "")
        after_text = edit.get("after_text", "")
        comment_text = edit.get("comment", "")
        tier = edit.get("tier", "")
        severity = edit.get("severity", "SUGGESTION")
        kata = edit.get("kata", "")

        if edit_type == "replace":
            success = rdoc.replace(original_text, new_text)
        elif edit_type == "delete":
            success = rdoc.delete(original_text)
        elif edit_type == "insert":
            success = rdoc.insert(after_text, new_text)
        else:
            # comment-only edit — no text matching needed
            rdoc.comment_count += 1
            success = True

        if not success and edit_type != "comment":
            logger.warning(
                f"Edit not applied — text not found: "
                f"[{tier}] '{original_text[:80]}...'"
            )

        if comment_text:
            target = original_text or after_text
            for i, para in enumerate(doc.paragraphs):
                if target and _normalize(target) in _normalize(para.text):
                    # Include kata in comment tag if present
                    if kata:
                        tag = f"[{tier}/{kata}]"
                    else:
                        tag = f"[{tier}]"
                    rdoc.comment(i, f"{tag} {comment_text}", severity)
                    break

    stem = input_path.stem
    redline_path = output_dir / f"{stem}_reviewed.docx"
    summary_path = output_dir / f"{stem}_review_summary.md"

    output_dir.mkdir(parents=True, exist_ok=True)
    rdoc.save(redline_path)
    summary_path.write_text(generate_summary(
        edits,
        summary,
        audience,
        match_rate=rdoc.match_rate,
        match_count=rdoc.match_count,
        miss_count=rdoc.miss_count,
        comment_count=rdoc.comment_count,
    ))

    return {
        "redline": redline_path,
        "summary": summary_path,
        "match_rate": rdoc.match_rate,
        "total_edits": len(edits),
        "edits_applied": rdoc.match_count,
        "edits_missed": rdoc.miss_count,
        "comment_only": rdoc.comment_count,
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python scripts/apply_redlines.py <input.docx> <review.json> [output_dir]")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    review_file = Path(sys.argv[2])
    output_dir = Path(sys.argv[3]) if len(sys.argv) > 3 else Path("output")

    if not input_file.exists():
        print(f"Error: {input_file} not found", file=sys.stderr)
        sys.exit(1)
    if not review_file.exists():
        print(f"Error: {review_file} not found", file=sys.stderr)
        sys.exit(1)

    print(f"Applying review to: {input_file}")
    outputs = apply_review(input_file, review_file, output_dir)

    print(f"\nRedlined doc: {outputs['redline']}")
    print(f"Review summary: {outputs['summary']}")
    print(f"Match rate: {outputs['match_rate']:.0%}")
    applied = outputs['edits_applied']
    comment_only = outputs.get('comment_only', 0)
    missed = outputs['edits_missed']
    total = outputs['total_edits']
    print(f"Edits: {applied} applied, {comment_only} comment-only, {missed} not found ({total} total)")
